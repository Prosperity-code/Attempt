"""报告渲染：Markdown / Word(docx) / PDF 三种格式。"""
from __future__ import annotations

from collections import OrderedDict
from pathlib import Path
from xml.sax.saxutils import escape

from docx import Document
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import ParagraphStyle
from reportlab.lib.units import mm
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.cidfonts import UnicodeCIDFont
from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer

# reportlab 内置的 CJK 字体（无需外部字体文件，纯 pip 即可支持中文）
pdfmetrics.registerFont(UnicodeCIDFont("STSong-Light"))
PDF_FONT = "STSong-Light"


def _fmt_stars(n: int) -> str:
    return f"{n:,}"


def _group_by_source(news: list[dict]) -> OrderedDict:
    grouped: OrderedDict = OrderedDict()
    for it in news:
        grouped.setdefault(it["source"], []).append(it)
    return grouped


def _esc(s) -> str:
    """PDF 段落中的 XML 转义。"""
    return escape(str(s))


# ---------- Markdown ----------

def render_markdown(today, projects: list[dict], news: list[dict], used_llm: bool) -> str:
    date_str = today.isoformat()
    lines = [
        f"# 📊 每日 AI 动态与 GitHub 热榜（{date_str}）",
        "",
        f"> 数据统计：GitHub 热榜 {len(projects)} 项 · AI 动态 {len(news)} 条"
        + (" · 中文摘要由 DeepSeek 生成" if used_llm else ""),
        "",
        "## 一、GitHub 开源项目热度榜 TOP10",
        "",
    ]
    for i, p in enumerate(projects, 1):
        intro = p.get("introduction") or p.get("description") or "（暂无简介）"
        usage = p.get("usage") or ""
        lines.append(f"### {i}. {p['name']}  ⭐ {_fmt_stars(p['stars'])}")
        lines.append(f"**语言**: {p['language']} ｜ **Stars**: {_fmt_stars(p['stars'])} ｜ **Forks**: {_fmt_stars(p['forks'])}")
        lines.append(f"**简介**: {intro}")
        if usage:
            lines.append(f"**用途**: {usage}")
        lines.append(f"🔗 [项目地址]({p['url']})")
        lines.append("")
    lines += ["## 二、昨日 AI 领域动态", ""]
    for source, items in _group_by_source(news).items():
        lines.append(f"### {source}")
        lines.append("")
        for it in items:
            pts = f"（👍 {it['points']}）" if it.get("points") else ""
            lines.append(f"- **[{it['title']}]({it['url']})**{pts}")
            if it.get("summary"):
                lines.append(f"  - 摘要：{it['summary']}")
        lines.append("")
    return "\n".join(lines)


# ---------- Word ----------

def render_docx(today, projects: list[dict], news: list[dict], used_llm: bool, path: Path) -> None:
    doc = Document()
    doc.add_heading(f"每日 AI 动态与 GitHub 热榜（{today.isoformat()}）", level=0)
    doc.add_paragraph(
        f"数据统计：GitHub 热榜 {len(projects)} 项 · AI 动态 {len(news)} 条"
        + (" · 中文摘要由 DeepSeek 生成" if used_llm else "")
    )
    doc.add_heading("一、GitHub 开源项目热度榜 TOP10", level=1)
    for i, p in enumerate(projects, 1):
        doc.add_heading(f"{i}. {p['name']}  ⭐ {_fmt_stars(p['stars'])}", level=2)
        para = doc.add_paragraph()
        para.add_run(f"语言：{p['language']} ｜ Stars：{_fmt_stars(p['stars'])} ｜ Forks：{_fmt_stars(p['forks'])}\n")
        para.add_run(f"简介：{p.get('introduction') or p.get('description') or '（暂无简介）'}\n")
        if p.get("usage"):
            para.add_run(f"用途：{p['usage']}\n")
        para.add_run(f"链接：{p['url']}")
    doc.add_heading("二、昨日 AI 领域动态", level=1)
    for source, items in _group_by_source(news).items():
        doc.add_heading(source, level=2)
        for it in items:
            para = doc.add_paragraph()
            r = para.add_run(it["title"])
            r.bold = True
            if it.get("points"):
                para.add_run(f"（👍 {it['points']}）")
            para.add_run("\n" + it["url"])
            if it.get("summary"):
                para.add_run(f"\n摘要：{it['summary']}")
    doc.save(str(path))


# ---------- PDF ----------

def render_pdf(today, projects: list[dict], news: list[dict], used_llm: bool, path: Path) -> None:
    styles = {
        "title": ParagraphStyle("title", fontName=PDF_FONT, fontSize=18, leading=26, spaceAfter=10),
        "h1": ParagraphStyle("h1", fontName=PDF_FONT, fontSize=15, leading=22, spaceBefore=12, spaceAfter=6),
        "h2": ParagraphStyle("h2", fontName=PDF_FONT, fontSize=12, leading=18, spaceBefore=8, spaceAfter=4),
        "body": ParagraphStyle("body", fontName=PDF_FONT, fontSize=10.5, leading=16, spaceAfter=4),
    }
    doc = SimpleDocTemplate(
        str(path),
        pagesize=A4,
        leftMargin=20 * mm, rightMargin=20 * mm,
        topMargin=18 * mm, bottomMargin=18 * mm,
        title=f"每日 AI 动态与 GitHub 热榜（{today.isoformat()}）",
    )
    story = [Paragraph(f"每日 AI 动态与 GitHub 热榜（{today.isoformat()}）", styles["title"])]
    story.append(
        Paragraph(
            f"数据统计：GitHub 热榜 {len(projects)} 项 · AI 动态 {len(news)} 条"
            + (" · 中文摘要由 DeepSeek 生成" if used_llm else ""),
            styles["body"],
        )
    )
    story.append(Spacer(1, 6))
    story.append(Paragraph("一、GitHub 开源项目热度榜 TOP10", styles["h1"]))
    for i, p in enumerate(projects, 1):
        story.append(Paragraph(f"{i}. {_esc(p['name'])}　⭐ {_fmt_stars(p['stars'])}", styles["h2"]))
        body = (
            f"语言：{_esc(p['language'])} ｜ Stars：{_fmt_stars(p['stars'])} ｜ Forks：{_fmt_stars(p['forks'])}<br/>"
            f"简介：{_esc(p.get('introduction') or p.get('description') or '（暂无简介）')}<br/>"
        )
        if p.get("usage"):
            body += f"用途：{_esc(p['usage'])}<br/>"
        body += f'链接：<link href="{_esc(p["url"])}" color="blue">{_esc(p["url"])}</link>'
        story.append(Paragraph(body, styles["body"]))
    story.append(Paragraph("二、昨日 AI 领域动态", styles["h1"]))
    for source, items in _group_by_source(news).items():
        story.append(Paragraph(_esc(source), styles["h2"]))
        for it in items:
            body = f"<b>{_esc(it['title'])}</b>" + (f"（👍 {it['points']}）" if it.get("points") else "")
            body += f'<br/><link href="{_esc(it["url"])}" color="blue">{_esc(it["url"])}</link>'
            if it.get("summary"):
                body += f"<br/>摘要：{_esc(it['summary'])}"
            story.append(Paragraph(body, styles["body"]))
    doc.build(story)
